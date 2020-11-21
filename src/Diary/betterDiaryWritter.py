import json
from docx import Document
from datetime import date, datetime
from src import config
Sections = {
    '-m': 'Main',
    '-w': 'Work',
    '-h': 'Happy',
    '-p': 'Plan'
}

def addContent(data, sections):
    """Allow user to input as much lines for each sections chosen"""
    for section in sections:
        # sections are the sections to add to in this session.
        print(section + ':')

        while True:
            sectionInput = input('> ')
            if sectionInput:
                data[section].append(sectionInput)
            else:
                break

    return data

def returnSections(command):
    """from the diary command work out what sections in particular to write, if no input, write all"""
    sections = []
    for shortcut in Sections:
        if shortcut in command:
            sections.append(Sections[shortcut])
    if not sections:
        sections = list(Sections.values()) #fill the list with values from dictionary (default fill all)
    return sections

def writeToDiary(data):
    """actually add the stuff in the cache to diary in a structured format"""
    try:
        diary = Document(config.DIARY_FILEPATH) #open the diary document object
    except:
        diary = Document()
    # Try-except handles the first time using the script, creating a new diary

    diary.add_paragraph(
        text = data["_Today"][0] + ' ' + data["_Today"][1] +
               '\nMain:' + ' '.join(data["Main"]) +
               '\nWork: ' + '\n-- '.join(data["Work"]) +
               '\nHappy Thoughts:' + ' '.join(data["Happy"]) +
               '\nPlans: ' + '\n-- '.join(data["Plan"])
    )
    diary.save(config.DIARY_FILEPATH)
    print('Finished writing diary for ' + data["_Today"][0] + ' , Hooray!')

def checkToday(data):
    """check if _Today section is already filled, or if it is time to inquire whether to log to diary now"""
    weekdays = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday', 'Sunday']
    today = date.today()

    if not data['_Today']:
        return [today.strftime('%Y/%m/%d'), weekdays[date.weekday(today)]]
        #return the values of today in the form of [ Y/M/D, weekday ], if the _Today section was previously blank

    elif datetime.now().hour > 22 or today.strftime('%Y/%m/%d') != data['_Today'][0]:
        return 1
        #return 1 if it is past 23:00 or the date has changed from the last diary cache. Value passed on to inquire if to log to diary

def clearDiaryCache(filepath):
    """empties diaryCache.json"""
    data = {
        'Main': [''],
        'Work': [''],
        'Happy': [''],
        'Plan': [''],
        '_Today': []
    }
    with open(filepath, 'w') as filewrite:
        json.dump(data, filewrite)

def main(filepath, command):
    with open(filepath, 'r') as fileread:
        data = json.load(fileread)
    #load the data from the diaryCache first

    check = checkToday(data) #use checkToday function to see if _Today was populated, and whether to remind writing to diary

    if 'write' in command:
        writeToDiary(data) #use writeToDiary function to write data from cache to diary in formatted form
        clearDiaryCache(filepath) #use clearDiaryCache to empty diaryCache.json
        return

    if check == 1:
        print(data["_Today"][0] + ' ' + data["_Today"][1] +
              '\nMain:' + ' '.join(data["Main"]) +
              '\nWork: ' + '\n-- '.join(data["Work"]) +
              '\nHappy Thoughts:' + ' '.join(data["Happy"]) +
              '\nPlans: ' + '\n-- '.join(data["Plan"]) + '\n')
        # show
        if input('write to diary now? (y/n)') == 'y':
            #if check condition met, ask if to write to diary
            writeToDiary(data) #use writeToDiary function to write data from cache to diary in formatted form
            clearDiaryCache(filepath) #use clearDiaryCache to empty diaryCache.json
            return
    elif check:
        #the _Today section was previously blank
        data['_Today'] = check

    print(data["_Today"][0] + ' ' + data["_Today"][1] +
          '\nMain:' + ' '.join(data["Main"]) +
          '\nWork: ' + '\n-- '.join(data["Work"]) +
          '\nHappy Thoughts:' + ' '.join(data["Happy"]) +
          '\nPlans: ' + '\n-- '.join(data["Plan"]) +'\n')

    sections = returnSections(command) #use returnSections function to find sections to write
    data = addContent(data, sections) #update content using addContent function

    with open(filepath, 'w') as filewrite:
        json.dump(data, filewrite)
    #save the data as json format and write into diaryCache

if __name__ == '__main__':
    main('diaryCache.json',' ')