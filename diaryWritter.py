#this module aims to help write diary in a formatted way.
#Note that this module was the first module of this program.

from docx import Document
from datetime import date
import os
import config

def main():
    # Start: open diary
    try:
        diary = Document(config.DIARY_FILEPATH)
    except:
        diary = Document()
    completed = False

    while not completed:
        # Date:
        weekdays = ['Monday','Tuesday','Wednesday','Thursday','Friday','Saturday','Sunday']
        today = date.today()

        # Main:
        # Work:
        # Plan: TODO: make this bullet point instead of text
        main = input('Main:')
        work = input('Work:')
        happy = input('Happy thoughts:')
        plan = input('Plan:')

        if input('Confirm(y/n):') == 'y':
            diary.add_paragraph(
                text=today.strftime('%Y/%m/%d') + ' ' + weekdays[date.weekday(today)] + ('\nMain: ' + main) + '\nWork: ' + work + '\nHappy thoughts: ' + happy +'\nPlan: ' + plan + '\n'
            )
            completed = True
            diary.save(config.DIARY_FILEPATH)
            os.system('"'+config.DIARY_FILEPATH+'"') #additional quotes to make it a single filepath in CMD
        elif input('Restart(y/n):') == 'n':
            completed = True

if __name__ == '__main__':
    main()